import streamlit as st
from openai import OpenAI
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import io
import random

st.set_page_config(page_title="SmartReport AI", layout="wide")

st.title("📄 SmartReport AI - Advanced Report Generator")

# OpenAI Client
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

# Sections
sections = [
    "Abstract", "Introduction", "Literature Review",
    "Methodology", "Results", "Conclusion", "References"
]

# UI Layout
col1, col2 = st.columns(2)

with col1:
    selected_sections = st.multiselect(
        "📌 Choose sections:",
        sections,
        default=sections
    )

with col2:
    shuffle = st.checkbox("🔀 Shuffle Sections")

ordered_sections = st.text_area(
    "✏️ Enter sections order (comma separated):",
    ",".join(selected_sections)
)

topic = st.text_input("🧠 Enter your project topic:")

# Generate
if st.button("🚀 Generate Report"):

    if not topic.strip():
        st.warning("⚠️ Enter topic")

    else:
        with st.spinner("Generating report... ⏳"):

            sections_list = [s.strip() for s in ordered_sections.split(",") if s.strip()]

            if shuffle:
                random.shuffle(sections_list)

            sections_text = ", ".join(sections_list)

            prompt = f"""
Write a detailed academic project report.

Topic: {topic}

Include sections:
{sections_text}

Rules:
- Use proper headings
- Formal academic writing
- Clear explanation
"""

            try:
                response = client.responses.create(
                    model="gpt-4o-mini",
                    input=prompt
                )

                content = response.output_text

                # Save to session (IMPORTANT for edit feature)
                st.session_state["report"] = content

            except Exception as e:
                st.error(f"❌ Error: {e}")

# -------- PREVIEW + EDIT -------- #

if "report" in st.session_state:

    st.subheader("📄 Preview Report")

    edited_content = st.text_area(
        "✏️ Edit your report before downloading:",
        value=st.session_state["report"],
        height=400
    )

    # -------- DOWNLOAD OPTIONS -------- #

    col1, col2, col3 = st.columns(3)

    # TXT
    with col1:
        st.download_button(
            "📥 TXT",
            edited_content,
            "report.txt"
        )

    # DOCX
    with col2:
        doc = Document()
        doc.add_heading("Project Report", 0)
        doc.add_paragraph(edited_content)

        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)

        st.download_button(
            "📄 Word",
            docx_buffer.getvalue(),
            "report.docx"
        )

    # PDF
    with col3:
        pdf_buffer = io.BytesIO()
        styles = getSampleStyleSheet()
        pdf = SimpleDocTemplate(pdf_buffer)

        flowables = []
        for line in edited_content.split("\n"):
            flowables.append(Paragraph(line, styles["Normal"]))

        pdf.build(flowables)

        st.download_button(
            "📑 PDF",
            pdf_buffer.getvalue(),
            "report.pdf"
        )

    # Share
    st.subheader("📋 Share / Copy")
    st.text_area("Copy your report:", edited_content, height=200)

    st.success("✅ You can edit, preview, and download your report!")
